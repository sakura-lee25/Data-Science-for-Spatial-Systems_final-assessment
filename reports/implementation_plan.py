"""Generate implementation plan as Word document."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path


def create_plan_docx(output_path: Path) -> None:
    doc = Document()

    # Styles
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    # Title
    title = doc.add_heading("UK Road Traffic Accident Spatial Analysis", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("Implementation Plan")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    # Context
    doc.add_heading("Context", level=1)
    doc.add_paragraph(
        "项目方向从原来的天气关联分析调整为空间统计 + MGWR + 机器学习分类。"
        "基于设计大纲，需要完成以下核心分析："
    )
    items = [
        "Global/Local Moran's I 空间自相关分析",
        "MGWR 多尺度地理加权回归",
        "Random Forest 不平衡分类",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Number")

    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("空间单元: ")
    r.bold = True
    p.add_run("全部 MSOA 级别")

    p = doc.add_paragraph()
    r = p.add_run("数据年份: ")
    r.bold = True
    p.add_run("2021-2024（4年）")

    bullets = [
        "2021 为疫情后首年恢复期，需在报告中注明",
        "主分析使用 2021-2024 合并数据",
        "敏感性分析：分别分析 2021/2022/2023/2024 各年空间格局，增强结论稳健性",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    # Phase 0
    doc.add_heading("Phase 0: 依赖更新与配置", level=1)
    p = doc.add_paragraph()
    r = p.add_run("文件: ")
    r.bold = True
    p.add_run("requirements.txt, .gitignore, src/utils/config.py")

    steps = [
        "更新 requirements.txt，添加：libpysal>=4.10.0, esda>=2.6.0 (空间自相关), mgwr>=2.2.1 (MGWR), imbalanced-learn>=0.12.0 (SMOTE)",
        "安装依赖: pip install -r requirements.txt",
        ".gitignore 添加 data/processed/*.gpkg, data/processed/*.pkl",
        "创建 src/utils/config.py — 集中管理路径和常量",
    ]
    for s in steps:
        doc.add_paragraph(s, style="List Number")

    # Phase 1
    doc.add_heading("Phase 1: 数据采集", level=1)
    p = doc.add_paragraph()
    r = p.add_run("目录: ")
    r.bold = True
    p.add_run("src/scraper/")
    p = doc.add_paragraph()
    r = p.add_run("输出: ")
    r.bold = True
    p.add_run("data/raw/ 中的原始文件")

    # Table for Phase 1
    table = doc.add_table(rows=7, cols=3)
    table.style = "Light Grid Accent 1"
    headers = ["文件", "数据源", "说明"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    data = [
        ("stats19.py", "STATS19 碰撞 CSV", "data.gov.uk, 年份 2021-2024"),
        ("boundaries.py", "MSOA 2021 边界 GeoJSON", "ONS Open Geography Portal"),
        ("imd.py", "IMD 2019 Excel", "MHCLG, LSOA 级别"),
        ("os_roads.py", "OS Open Roads GeoPackage", "OS Data Hub"),
        ("population.py", "MSOA 人口估计", "ONS mid-year estimates"),
        ("main.py", "编排脚本", "顺序调用上述所有下载器"),
    ]
    for i, (a, b, c) in enumerate(data):
        table.rows[i + 1].cells[0].text = a
        table.rows[i + 1].cells[1].text = b
        table.rows[i + 1].cells[2].text = c

    doc.add_paragraph()
    doc.add_heading("关键下载 URL", level=2)
    urls = [
        "STATS19: https://data.dft.gov.uk/road-accidents-safety-data/dft-road-casualty-statistics-collision-{year}.csv",
        "MSOA: ONS WFS endpoint (EPSG:27700)",
        "IMD: https://assets.publishing.service.gov.uk/.../File_7_-_All_IoD2019_Scores.xlsx",
        "需同时下载 LSOA→MSOA 查找表",
    ]
    for u in urls:
        doc.add_paragraph(u, style="List Bullet")

    p = doc.add_paragraph()
    r = p.add_run("验证: ")
    r.bold = True
    p.add_run("python -m src.scraper.main 成功运行，data/raw/ 包含所有文件")

    # Phase 2
    doc.add_heading("Phase 2: 数据处理与特征工程", level=1)
    p = doc.add_paragraph()
    r = p.add_run("目录: ")
    r.bold = True
    p.add_run("src/analysis/preprocessing/, src/utils/")
    p = doc.add_paragraph()
    r = p.add_run("输出: ")
    r.bold = True
    p.add_run("data/processed/msoa_analysis.gpkg")

    table2 = doc.add_table(rows=9, cols=2)
    table2.style = "Light Grid Accent 1"
    table2.rows[0].cells[0].text = "文件"
    table2.rows[0].cells[1].text = "功能"
    data2 = [
        ("src/utils/config.py", "路径配置、常量"),
        ("src/utils/io.py", "分块读取大 CSV"),
        ("src/utils/crs.py", "OSGB36↔WGS84 转换"),
        ("preprocessing/stats19_cleaner.py", "清洗 STATS19 数据"),
        ("preprocessing/spatial_join.py", "事故点→MSOA 空间联结 + 聚合"),
        ("preprocessing/imd_aggregator.py", "IMD LSOA→MSOA 聚合"),
        ("preprocessing/road_metrics.py", "道路长度 + 路口密度计算"),
        ("preprocessing/feature_builder.py", "合并所有特征为单一 GeoDataFrame"),
    ]
    for i, (a, b) in enumerate(data2):
        table2.rows[i + 1].cells[0].text = a
        table2.rows[i + 1].cells[1].text = b

    doc.add_paragraph()
    doc.add_heading("核心输出字段", level=2)
    fields = [
        "accident_rate_per_10k = accident_count / population × 10000",
        "imd_score, road_length_km, road_density, junction_density",
        "urban_pct, wet_road_pct, dark_pct",
    ]
    for f in fields:
        doc.add_paragraph(f, style="List Bullet")

    doc.add_heading("敏感性分析数据", level=2)
    doc.add_paragraph("data/processed/msoa_analysis.gpkg — 2021-2024 合并主数据", style="List Bullet")
    doc.add_paragraph("data/processed/msoa_yearly_{year}.gpkg — 各年单独数据 (2021/2022/2023/2024)", style="List Bullet")

    p = doc.add_paragraph()
    r = p.add_run("验证: ")
    r.bold = True
    p.add_run("~6,791 行 (England MSOAs)，无 NaN，CRS=EPSG:27700")

    # Phase 3
    doc.add_heading("Phase 3: 空间自相关分析", level=1)
    p = doc.add_paragraph()
    r = p.add_run("文件: ")
    r.bold = True
    p.add_run("src/analysis/spatial_autocorr.py, notebooks/01_spatial_autocorrelation.ipynb")

    steps3 = [
        "构建 Queen 邻接空间权重矩阵 (libpysal)",
        "Global Moran's I 检验整体空间聚集",
        "Local Moran's I (LISA) 识别 HH/LL/HL/LH 聚类",
        "生成 Moran 散点图 + LISA 聚类地图",
        "敏感性分析: 分别对 2021/2022/2023/2024 各年运行 Global Moran's I + LISA，对比空间格局稳定性（注明 2021 为疫情恢复期）",
    ]
    for s in steps3:
        doc.add_paragraph(s, style="List Number")

    doc.add_heading("输出图表", level=2)
    figs3 = [
        "fig01_accident_rate_choropleth.png",
        "fig02_moran_scatter.png",
        "fig03_lisa_clusters.png",
        "fig03b_lisa_yearly_comparison.png (4年对比面板图)",
    ]
    for f in figs3:
        doc.add_paragraph(f, style="List Bullet")

    # Phase 4
    doc.add_heading("Phase 4: MGWR 分析", level=1)
    p = doc.add_paragraph()
    r = p.add_run("文件: ")
    r.bold = True
    p.add_run("src/analysis/mgwr_analysis.py, notebooks/02_mgwr_analysis.ipynb")

    steps4 = [
        "VIF 多重共线性检查",
        "OLS 基线回归 → 残差 Moran's I (证明需要 GWR)",
        "MGWR 拟合 (自适应核, AICc 带宽选择)",
        "提取局部系数面 → 保存 mgwr_coefficients.gpkg",
        "生成系数空间分布图",
    ]
    for s in steps4:
        doc.add_paragraph(s, style="List Number")

    p = doc.add_paragraph()
    r = p.add_run("模型变量: ")
    r.bold = True
    p.add_run("log_accident_rate ~ imd_score + junction_density + road_density + urban_pct + dark_pct + wet_road_pct")

    p = doc.add_paragraph()
    r = p.add_run("⚠ 注意: ")
    r.bold = True
    r.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    p.add_run("MGWR 计算约需 2-4 小时，必须缓存结果")

    # Phase 5
    doc.add_heading("Phase 5: 机器学习分类", level=1)
    p = doc.add_paragraph()
    r = p.add_run("文件: ")
    r.bold = True
    p.add_run("src/analysis/ml_classification.py, notebooks/03_ml_classification.ipynb")

    steps5 = [
        "定义高风险标签 (top 20% 事故率 → label=1)",
        "空间感知 train/test split (空间分块避免泄漏)",
        "Random Forest + SMOTE",
        "阈值优化 (交叉验证)",
        "评估: PR 曲线、混淆矩阵、特征重要性",
    ]
    for s in steps5:
        doc.add_paragraph(s, style="List Number")

    doc.add_heading("输出图表", level=2)
    figs5 = [
        "fig07_pr_curve.png",
        "fig08_confusion_matrix.png",
        "fig09_feature_importance.png",
    ]
    for f in figs5:
        doc.add_paragraph(f, style="List Bullet")

    # Phase 6
    doc.add_heading("Phase 6: 可视化与报告", level=1)
    p = doc.add_paragraph()
    r = p.add_run("目录: ")
    r.bold = True
    p.add_run("src/visualization/")

    table6 = doc.add_table(rows=3, cols=2)
    table6.style = "Light Grid Accent 1"
    table6.rows[0].cells[0].text = "文件"
    table6.rows[0].cells[1].text = "功能"
    table6.rows[1].cells[0].text = "maps.py"
    table6.rows[1].cells[1].text = "Choropleth、LISA、MGWR 系数地图"
    table6.rows[2].cells[0].text = "statistical_plots.py"
    table6.rows[2].cells[1].text = "Moran 散点图、PR 曲线、混淆矩阵"

    doc.add_paragraph()
    doc.add_paragraph("所有图表 300 DPI，学术风格 (Serif 字体，白色背景)")

    # Dependency
    doc.add_heading("依赖关系与执行顺序", level=1)
    doc.add_paragraph("Phase 0 → Phase 1 → Phase 2 → Phase 3/4/5 (可并行) → Phase 6")
    doc.add_paragraph("按 Phase 0 → 1 → 2 → 3 → 4 → 5 → 6 逐步实施。每完成一个 Phase，验证后再进入下一个。")

    doc.save(str(output_path))


if __name__ == "__main__":
    output = Path(__file__).parent / "implementation_plan.docx"
    create_plan_docx(output)
    print(f"Word document saved to: {output}")
